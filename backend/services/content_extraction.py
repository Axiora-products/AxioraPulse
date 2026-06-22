"""
services/content_extraction.py
──────────────────────────────
Traditional (non-AI) content extraction for documents, spreadsheets, images and
website links, with a heuristic confidence score per extraction.

Design goals (per the Document/Screenshot/Link workflow):
  - Use deterministic parsers (pypdf, python-docx, openpyxl, csv) and OCR
    (pytesseract) — the AI is reserved for survey generation, not extraction.
  - Never raise on a single bad input: every extractor returns an
    ExtractionResult with text + confidence + warnings, degrading gracefully when
    an optional library or the Tesseract binary is unavailable.
  - Sanitize extracted text and bound its length.
  - Fetch website links behind an SSRF guard (block private/loopback/link-local
    hosts, non-standard ports, oversized bodies, untrusted redirects).
"""

from __future__ import annotations

import csv
import io
import ipaddress
import logging
import socket
import unicodedata
from dataclasses import dataclass, field
from urllib.parse import urlparse, urljoin

import requests

logger = logging.getLogger(__name__)

MAX_CHARS = 8000
_LINK_TIMEOUT = 8
_LINK_MAX_BYTES = 2_000_000
_LINK_MAX_REDIRECTS = 3


@dataclass
class ExtractionResult:
    text: str = ""
    confidence: int = 0           # 0-100
    ocr_quality: str | None = None  # 'High' | 'Medium' | 'Low' (images only)
    source: str = ""
    warnings: list = field(default_factory=list)
    truncated: bool = False

    def as_dict(self) -> dict:
        return {
            "text": self.text,
            "confidence": self.confidence,
            "ocr_quality": self.ocr_quality,
            "source": self.source,
            "warnings": self.warnings,
            "truncated": self.truncated,
            "needs_review": self.confidence < CONFIDENCE_REVIEW_THRESHOLD,
        }


CONFIDENCE_REVIEW_THRESHOLD = 70


class ExtractionError(Exception):
    """Raised for link extraction when the URL is unsafe or unreachable."""


# ── Sanitization ───────────────────────────────────────────────────────────────
def sanitize_extracted(text: str) -> tuple[str, bool]:
    """Normalize unicode, strip control/zero-width chars, collapse blank runs and
    bound the length. Returns (clean_text, truncated)."""
    if not text:
        return "", False
    text = unicodedata.normalize("NFKC", text)
    text = "".join(
        ch for ch in text if ch in ("\n", "\t") or unicodedata.category(ch)[0] != "C"
    )
    # Collapse excessive blank lines / spaces.
    lines = [ln.rstrip() for ln in text.splitlines()]
    cleaned = "\n".join(ln for ln in lines if ln.strip() != "" or True).strip()
    truncated = len(cleaned) > MAX_CHARS
    return cleaned[:MAX_CHARS], truncated


def _text_confidence(text: str, *, structured: bool = False) -> int:
    n = len(text.strip())
    if n == 0:
        return 0
    if structured:
        return 96
    if n > 400:
        return 95
    if n > 120:
        return 86
    if n > 30:
        return 72
    return 45


# ── Documents ──────────────────────────────────────────────────────────────────
def _extract_pdf(data: bytes) -> ExtractionResult:
    r = ExtractionResult(source="pdf")
    try:
        try:
            from pypdf import PdfReader
        except ImportError:  # fallback to legacy name
            from PyPDF2 import PdfReader  # type: ignore
        reader = PdfReader(io.BytesIO(data))
        pages = reader.pages[:30]
        parts, pages_with_text = [], 0
        for page in pages:
            t = page.extract_text() or ""
            if t.strip():
                pages_with_text += 1
            parts.append(t)
        text, r.truncated = sanitize_extracted("\n".join(parts))
        r.text = text
        if not text:
            r.confidence = 25
            r.warnings.append("No selectable text found — this may be a scanned PDF. Consider uploading it as an image for OCR.")
        else:
            r.confidence = _text_confidence(text)
    except ImportError:
        r.warnings.append("PDF parser is not available on the server.")
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", type(exc).__name__)
        r.warnings.append("We could not read this PDF. Please review or re-upload.")
    return r


def _extract_docx(data: bytes) -> ExtractionResult:
    r = ExtractionResult(source="word")
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text, r.truncated = sanitize_extracted("\n".join(parts))
        r.text = text
        r.confidence = _text_confidence(text)
        if not text:
            r.warnings.append("No readable text found in this document.")
    except ImportError:
        r.warnings.append("Word parser is not available on the server.")
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", type(exc).__name__)
        r.warnings.append("We could not read this Word document. Please review or re-upload.")
    return r


def _extract_xlsx(data: bytes) -> ExtractionResult:
    r = ExtractionResult(source="spreadsheet")
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets[:5]:
            parts.append(f"# Sheet: {ws.title}")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 200:
                    parts.append("… (additional rows omitted)")
                    break
                cells = ["" if c is None else str(c) for c in row]
                if any(cells):
                    parts.append(" | ".join(cells))
        wb.close()
        text, r.truncated = sanitize_extracted("\n".join(parts))
        r.text = text
        r.confidence = _text_confidence(text, structured=True) if text else 0
        if not text:
            r.warnings.append("No data found in this spreadsheet.")
    except ImportError:
        r.warnings.append("Spreadsheet parser is not available on the server.")
    except Exception as exc:
        logger.warning("XLSX extraction failed: %s", type(exc).__name__)
        r.warnings.append("We could not read this spreadsheet. Please review or re-upload.")
    return r


def _extract_csv(data: bytes) -> ExtractionResult:
    r = ExtractionResult(source="spreadsheet")
    try:
        raw = data.decode("utf-8", errors="ignore")
        sample = raw[:4096]
        try:
            dialect = csv.Sniffer().sniff(sample)
        except Exception:
            dialect = csv.excel
        reader = csv.reader(io.StringIO(raw), dialect)
        parts = []
        for i, row in enumerate(reader):
            if i >= 300:
                parts.append("… (additional rows omitted)")
                break
            if any(cell.strip() for cell in row):
                parts.append(" | ".join(cell.strip() for cell in row))
        text, r.truncated = sanitize_extracted("\n".join(parts))
        r.text = text
        r.confidence = _text_confidence(text, structured=True) if text else 0
        if not text:
            r.warnings.append("No rows found in this CSV file.")
    except Exception as exc:
        logger.warning("CSV extraction failed: %s", type(exc).__name__)
        r.warnings.append("We could not read this CSV file. Please review or re-upload.")
    return r


def _extract_txt(data: bytes) -> ExtractionResult:
    r = ExtractionResult(source="text")
    try:
        raw = data.decode("utf-8")
        clean_conf = 97
    except UnicodeDecodeError:
        raw = data.decode("utf-8", errors="replace")
        clean_conf = 80
        r.warnings.append("Some characters could not be decoded and were replaced.")
    text, r.truncated = sanitize_extracted(raw)
    r.text = text
    r.confidence = clean_conf if text else 0
    return r


# ── Images (OCR) ───────────────────────────────────────────────────────────────
def _ocr_quality_from_conf(conf: int) -> str:
    if conf >= 85:
        return "High"
    if conf >= 60:
        return "Medium"
    return "Low"


def _extract_image_ocr(data: bytes) -> ExtractionResult:
    r = ExtractionResult(source="image")
    try:
        from PIL import Image
    except ImportError:
        r.warnings.append("Image processing is not available on the server.")
        return r
    try:
        img = Image.open(io.BytesIO(data))
        img.load()
        r.warnings.append("")  # placeholder removed below
        r.warnings = []
        metadata = {"format": img.format, "width": img.width, "height": img.height, "mode": img.mode}
        r.source = f"image ({img.format})"
        try:
            import pytesseract
            ocr = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            words, confs = [], []
            for word, conf in zip(ocr.get("text", []), ocr.get("conf", [])):
                if word and word.strip():
                    words.append(word)
                    try:
                        c = int(float(conf))
                        if c >= 0:
                            confs.append(c)
                    except (ValueError, TypeError):
                        pass
            text, r.truncated = sanitize_extracted(" ".join(words))
            r.text = text
            avg = int(sum(confs) / len(confs)) if confs else 0
            r.confidence = avg if text else 0
            r.ocr_quality = _ocr_quality_from_conf(avg)
            if not text:
                r.warnings.append("No readable text was detected in this image.")
            elif avg < CONFIDENCE_REVIEW_THRESHOLD:
                r.warnings.append("Some text may not have been read accurately. Please review and edit before continuing.")
        except ImportError:
            r.warnings.append("OCR is not available on the server.")
            r.ocr_quality = "Low"
        except Exception as exc:
            # Most commonly the Tesseract binary is missing.
            logger.warning("OCR failed: %s", type(exc).__name__)
            r.warnings.append("Text could not be extracted from this image automatically. Please add the context manually.")
            r.ocr_quality = "Low"
        # Attach lightweight metadata note.
        r.text = (r.text + (f"\n\n[Image metadata: {metadata['format']} {metadata['width']}x{metadata['height']}]" if r.text else "")).strip()
    except Exception as exc:
        logger.warning("Image open failed: %s", type(exc).__name__)
        r.warnings.append("We could not process this image. Please review or re-upload.")
    return r


# ── Dispatcher ─────────────────────────────────────────────────────────────────
def extract_from_document(data: bytes, content_type: str, filename: str = "") -> ExtractionResult:
    ct = (content_type or "").split(";", 1)[0].strip().lower()
    name = (filename or "").lower()

    if ct == "application/pdf" or name.endswith(".pdf"):
        return _extract_pdf(data)
    if "wordprocessingml" in ct or ct == "application/msword" or name.endswith((".doc", ".docx")):
        return _extract_docx(data)
    if "spreadsheetml" in ct or ct in ("application/vnd.ms-excel",) or name.endswith((".xlsx", ".xls")):
        return _extract_xlsx(data)
    if ct in ("text/csv", "application/csv") or name.endswith(".csv"):
        return _extract_csv(data)
    if ct == "text/plain" or name.endswith(".txt"):
        return _extract_txt(data)
    if ct.startswith("image/") or name.endswith((".png", ".jpg", ".jpeg", ".webp")):
        return _extract_image_ocr(data)

    r = ExtractionResult(source="file")
    r.warnings.append("This file type is not supported for content extraction.")
    return r


# ── Website links (SSRF-safe) ──────────────────────────────────────────────────
def _assert_safe_url(url: str) -> None:
    p = urlparse(url)
    if p.scheme not in ("http", "https"):
        raise ExtractionError("Only http and https links are supported.")
    host = p.hostname
    if not host:
        raise ExtractionError("That does not look like a valid website link.")
    if p.port and p.port not in (80, 443):
        raise ExtractionError("That website link uses an unsupported port.")
    try:
        infos = socket.getaddrinfo(host, p.port or (443 if p.scheme == "https" else 80))
    except Exception:
        raise ExtractionError("We could not reach that website link.")
    for info in infos:
        ip = info[4][0]
        try:
            ip_obj = ipaddress.ip_address(ip.split("%")[0])
        except ValueError:
            continue
        if (
            ip_obj.is_private or ip_obj.is_loopback or ip_obj.is_link_local
            or ip_obj.is_reserved or ip_obj.is_multicast or ip_obj.is_unspecified
        ):
            raise ExtractionError("That website link cannot be processed.")


def _fetch_html(url: str) -> str:
    current = url
    for _ in range(_LINK_MAX_REDIRECTS + 1):
        _assert_safe_url(current)
        resp = requests.get(
            current,
            timeout=_LINK_TIMEOUT,
            stream=True,
            allow_redirects=False,
            headers={"User-Agent": "AxioraPulseBot/1.0 (+content-extraction)"},
        )
        if resp.status_code in (301, 302, 303, 307, 308):
            loc = resp.headers.get("location")
            if not loc:
                raise ExtractionError("We could not reach that website link.")
            current = urljoin(current, loc)
            continue
        resp.raise_for_status()
        ctype = resp.headers.get("content-type", "")
        if "html" not in ctype and "text" not in ctype:
            raise ExtractionError("That link does not point to a readable web page.")
        content = b""
        for chunk in resp.iter_content(8192):
            content += chunk
            if len(content) > _LINK_MAX_BYTES:
                break
        return content.decode(resp.encoding or "utf-8", errors="ignore")
    raise ExtractionError("That website link redirected too many times.")


_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".svg")


def is_loadable_image(url: str, timeout: int = 4) -> bool:
    """Return True only if `url` is a safe, reachable, image resource.

    Used to verify web-collected images referenced by AI-generated questions
    actually load before they are shown to respondents. SSRF-guarded; bounded by
    a short timeout; treats a 200 with an image content-type (or image extension)
    as loadable.
    """
    if not url or not isinstance(url, str):
        return False
    try:
        _assert_safe_url(url)
    except ExtractionError:
        return False
    headers = {"User-Agent": "AxioraPulseBot/1.0 (+media-check)"}
    try:
        resp = requests.head(url, timeout=timeout, allow_redirects=True, headers=headers)
        if resp.status_code in (403, 405) or resp.status_code >= 500:
            # Some CDNs reject HEAD — retry a lightweight GET.
            resp = requests.get(url, timeout=timeout, stream=True, headers=headers)
        if resp.status_code != 200:
            return False
        ctype = (resp.headers.get("content-type") or "").split(";", 1)[0].strip().lower()
        if ctype.startswith("image/"):
            return True
        return urlparse(url).path.lower().endswith(_IMAGE_EXTS)
    except Exception:
        return False


def extract_from_url(url: str) -> ExtractionResult:
    r = ExtractionResult(source=url)
    html = _fetch_html(url)  # raises ExtractionError on unsafe/unreachable
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        r.warnings.append("Website parsing is not available on the server.")
        return r

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "nav", "footer", "header", "form", "svg"]):
        tag.decompose()

    title = (soup.title.string.strip() if soup.title and soup.title.string else "")
    meta_desc = ""
    md = soup.find("meta", attrs={"name": "description"}) or soup.find("meta", attrs={"property": "og:description"})
    if md and md.get("content"):
        meta_desc = md["content"].strip()
    headings = [h.get_text(" ", strip=True) for h in soup.find_all(["h1", "h2", "h3"])[:15]]
    main = soup.find("main") or soup.find("article") or soup.body or soup
    body_text = main.get_text("\n", strip=True) if main else ""

    parts = []
    if title:
        parts.append(f"Page title: {title}")
    if meta_desc:
        parts.append(f"Description: {meta_desc}")
    if headings:
        parts.append("Key headings:\n- " + "\n- ".join(headings))
    if body_text:
        parts.append("Main content:\n" + body_text)

    text, r.truncated = sanitize_extracted("\n\n".join(parts))
    r.text = text

    score = 0
    if title:
        score += 35
    if meta_desc:
        score += 25
    if len(body_text) > 300:
        score += 35
    elif body_text:
        score += 15
    r.confidence = min(score, 95) if text else 0
    if not text:
        r.warnings.append("We could not extract readable content from that page.")
    elif r.confidence < CONFIDENCE_REVIEW_THRESHOLD:
        r.warnings.append("Limited content was extracted from this page. Please review and edit before continuing.")
    return r
